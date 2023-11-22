import os
from docx import Document
from dotenv import load_dotenv
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Mm


def set_table_style(table):
    # 表のスタイルを設定
    table.style = "Table Grid"
    # 表の位置を中央に設定
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(table.rows):
        # 表の高さを設定
        if i == 0:
            row.height = Mm(6)
        else:
            row.height = Mm(13)
        # 表の幅を設定
        for j, cell in enumerate(row.cells):
            if j == 0:
                cell.width = Mm(18)
            elif j == 3 or j == 4:
                cell.width = Mm(20)
            else:
                cell.width = Mm(23)
    return table


def main():
    user = os.getenv("USERNAME")
    file = Document()

    # ファイルの余白を設定
    section = file.sections[0]
    section.top_margin = Mm(30)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(20)
    section.right_margin = Mm(16)

    paragraph_format = file.styles["Normal"].paragraph_format
    # デフォルトの段落スペースを0ポイントに設定
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)

    p = file.add_paragraph("報告日：2023/xx/xx")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.right_indent = Pt(30)

    file.add_paragraph()
    p = file.add_paragraph("学籍番号：xxxxxx")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.right_indent = Pt(75)
    p = file.add_paragraph("氏名：xxxxxx")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.right_indent = Pt(75)

    p = file.add_paragraph("2023年度 週間進捗報告")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    file.add_paragraph()
    file.add_paragraph("【2023年 11月第x週】")

    table = file.add_table(rows=3, cols=8)
    table = set_table_style(table)

    for i in range(3):
        for j in range(8):
            # 各セルに入れる文字を設定
            if i == 1 and j == 0:
                table.cell(i, j).text = "ゼミ室\n入退"
            if i == 2 and j == 0:
                table.cell(i, j).text = "備考"

    file.add_paragraph()
    file.add_paragraph("・予定と結果(概要)")
    file.add_paragraph(" ・")
    file.add_paragraph()
    file.add_paragraph("・予定していたこと以外で行ったこと(概要)")
    file.add_paragraph()
    file.add_paragraph("・その他（先輩、先生からの指示、同輩等からの助言、他の人に伝えたい技術、ニュースなど）")
    file.add_paragraph()
    file.add_paragraph("・直近1週間の振り返り（自己評価、今後に向けた課題）")
    file.add_paragraph()

    file.save("/mnt/c/Users/" + user + "/Downloads/sample.docx")


if __name__ == "__main__":
    load_dotenv()
    main()
